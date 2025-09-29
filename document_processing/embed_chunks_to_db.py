import sys
import os
import json
import uuid
import asyncpg
import numpy as np
import PyPDF2
from pathlib import Path
from docx import Document
import asyncio

from chonkie import SemanticChunker
from sentence_transformers import SentenceTransformer

from typing import List, Dict, Any, Optional
from dataclasses import dataclass

# Import your existing chunking functions
sys.path.append(os.path.dirname(os.path.abspath(__file__)))


@dataclass
class Chunk:
    """Chunk data structure to match your existing interface."""
    id: str
    document_id: str
    text: str
    embedding: List[float]
    metadata: Optional[Dict] = None


class EmbeddingGenerator:
    """Generate embeddings using SentenceTransformers."""

    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        """
        Initialize embedding generator.

        Args:
            model_name: SentenceTransformer model name
        """
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)
        self.embedding_dim = self.model.get_sentence_embedding_dimension()

    def embed_text(self, text: str) -> List[float]:
        """
        Generate embedding for a single text.
        Args:
            text: Input text
        Returns:
            List of embedding values
        """
        embedding = self.model.encode(text)
        return embedding.tolist()

    def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings for multiple texts.
        Args:
            texts: List of input texts
        Returns:
            List of embedding lists
        """
        embeddings = self.model.encode(texts)
        return [emb.tolist() for emb in embeddings]


class VectorStore:
    """Vector store using pgvector for efficient similarity search."""

    def __init__(self, connection_params: Dict[str, str], table_name: str = "chunks"):
        """
        Initialize vector store with pgvector support.
        Args:
            connection_params: Database connection parameters
            table_name: Name of the chunks table
        """
        self.connection_params = connection_params
        self.table_name = table_name
        self.connection_string = self._build_connection_string()
        self._initialized = False

    def _build_connection_string(self) -> str:
        """Build asyncpg connection string from parameters."""
        return f"postgresql://{self.connection_params['user']}:{self.connection_params['password']}@{self.connection_params['host']}:{self.connection_params['port']}/{self.connection_params['dbname']}"

    async def _get_connection(self):
        """Get database connection with pgvector support."""
        conn = await asyncpg.connect(self.connection_string)
        # Enable pgvector extension
        await conn.execute("CREATE EXTENSION IF NOT EXISTS vector;")
        return conn

    async def _initialize_database(self):
        """Initialize database with pgvector extension and table."""
        try:
            if self._initialized:
                return

            conn = await self._get_connection()

            # Create table with proper vector column
            # Assuming 384-dimensional embeddings for all-MiniLM-L6-v2
            # Adjust dimension based on your model
            await conn.execute(f"""
            CREATE TABLE IF NOT EXISTS {self.table_name} (
                id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL,
                text TEXT NOT NULL,
                embedding vector(384),  -- Adjust dimension as needed
                metadata JSONB,
                created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
            );
            """)

            # Create index for similarity search
            await conn.execute(f"""
            CREATE INDEX IF NOT EXISTS {self.table_name}_embedding_idx
            ON {self.table_name} USING ivfflat (embedding vector_cosine_ops)
            WITH (lists = 100);
            """)

            # Create index on document_id for filtering
            await conn.execute(f"""
            CREATE INDEX IF NOT EXISTS {self.table_name}_document_id_idx
            ON {self.table_name} (document_id);
            """)

            await conn.close()
            self._initialized = True
            print(f"Database initialized with table: {self.table_name}")

        except Exception as e:
            print(f"Error initializing database: {e}")
            raise

    async def add_chunks(self, chunks: List[Chunk], batch_size: int = 100):
        """Add chunks to vector store using batch insert for efficiency."""
        try:
            if not self._initialized:
                await self._initialize_database()

            conn = await self._get_connection()

            # Prepare data for batch insert
            chunk_data = []
            for chunk in chunks:
                chunk_data.append((
                    chunk.id,
                    chunk.document_id,
                    chunk.text,
                    # Convert to list for asyncpg
                    chunk.embedding,
                    json.dumps(
                        chunk.metadata) if chunk.metadata else json.dumps({})
                ))

            # Use asyncpg's executemany for efficient batch insert
            insert_sql = f"""
            INSERT INTO {self.table_name} (id, document_id, text, embedding, metadata)
            VALUES ($1, $2, $3, $4, $5)
            ON CONFLICT (id) DO UPDATE SET
                document_id = EXCLUDED.document_id,
                text = EXCLUDED.text,
                embedding = EXCLUDED.embedding,
                metadata = EXCLUDED.metadata;
            """

            # Process in batches
            for i in range(0, len(chunk_data), batch_size):
                batch = chunk_data[i:i + batch_size]
                await conn.executemany(insert_sql, batch)

            await conn.close()
            print(f"Added {len(chunks)} chunks to vector store")

        except Exception as e:
            print(f"Error adding chunks: {e}")
            raise

    async def search_similar_chunks(self, query_embedding: List[float],
                              limit: int = 5, threshold: float = 0.7,
                              document_ids: Optional[List[str]] = None) -> List[Dict]:
        """
        Search for similar chunks using pgvector cosine similarity.

        Args:
            query_embedding: Query embedding vector
            limit: Maximum number of results
            threshold: Similarity threshold (0-1, higher = more similar)
            document_ids: Optional list of document IDs to filter by

        Returns:
            List of similar chunks with metadata
        """
        try:
            if not self._initialized:
                await self._initialize_database()

            conn = await self._get_connection()

            # Build query with optional document filtering
            base_query = f"""
                SELECT
                    id,
                    text,
                    metadata,
                    document_id,
                    (1 - (embedding <=> $1)) as similarity
                FROM {self.table_name}
                WHERE (1 - (embedding <=> $1)) >= $2
            """

            params = [query_embedding, threshold]

            if document_ids:
                base_query += " AND document_id = ANY($3)"
                params.append(document_ids)
                base_query += """
                    ORDER BY embedding <=> $1
                    LIMIT $4
                """
                params.append(limit)
            else:
                base_query += """
                    ORDER BY embedding <=> $1
                    LIMIT $3
                """
                params.append(limit)

            rows = await conn.fetch(base_query, *params)

            results = []
            for row in rows:
                results.append({
                    'chunk_id': row['id'],
                    'text': row['text'],
                    'metadata': row['metadata'] if isinstance(row['metadata'], (dict, type(None))) else json.loads(row['metadata']),
                    'document_id': row['document_id'],
                    'similarity': float(row['similarity'])
                })

            await conn.close()
            return results

        except Exception as e:
            print(f"Error searching chunks: {e}")
            raise

    async def delete_document_chunks(self, document_id: str) -> int:
        """Delete all chunks for a specific document."""
        try:
            if not self._initialized:
                await self._initialize_database()

            conn = await self._get_connection()
            result = await conn.execute(
                f"DELETE FROM {self.table_name} WHERE document_id = $1", document_id)
            deleted_count = int(result.split()[-1]) if result else 0
            await conn.close()
            print(
                f"Deleted {deleted_count} chunks for document: {document_id}")
            return deleted_count
        except Exception as e:
            print(f"Error deleting document chunks: {e}")
            raise

    async def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store."""
        try:
            if not self._initialized:
                await self._initialize_database()

            conn = await self._get_connection()
            row = await conn.fetchrow(f"""
            SELECT
                COUNT(*) as total_chunks,
                COUNT(DISTINCT document_id) as total_documents,
                AVG(LENGTH(text)) as avg_text_length,
                MIN(created_at) as earliest_chunk,
                MAX(created_at) as latest_chunk
            FROM {self.table_name}
            """)

            stats = {
                'total_chunks': row['total_chunks'],
                'total_documents': row['total_documents'],
                'avg_text_length': float(row['avg_text_length']) if row['avg_text_length'] else 0,
                'earliest_chunk': row['earliest_chunk'].isoformat() if row['earliest_chunk'] else None,
                'latest_chunk': row['latest_chunk'].isoformat() if row['latest_chunk'] else None
            }

            await conn.close()
            return stats
        except Exception as e:
            print(f"Error getting stats: {e}")
            raise


class ChunkEmbeddingPipeline:
    """Complete pipeline for chunking documents and storing embeddings."""

    def __init__(self, db_params: Dict[str, str], embedding_model: str,
                 table_name: str):
        """
        Initialize the pipeline.
        Args:
            db_params: Database connection parameters
            embedding_model: SentenceTransformer model name
            table_name: Name of the chunks table
        """
        self.embedding_generator = EmbeddingGenerator(embedding_model)
        self.vector_store = VectorStore(db_params, table_name)

    def extract_text_from_pdf(self, pdf_path: str) -> tuple:
        """Extract text from PDF file with page tracking."""
        text = ""
        page_mapping = []

        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                start_pos = len(text)
                text += page_text + "\n"
                end_pos = len(text) - 1

                if page_text.strip():
                    page_mapping.append((start_pos, end_pos, page_num + 1))

        return text, page_mapping

    def extract_text_from_docx(self, docx_path: str) -> tuple:
        """Extract text from DOCX file with page estimation."""
        try:
            from docx import Document

            doc = Document(docx_path)
            text = ""
            page_mapping = []
            chars_per_page = 2500  # Rough estimate for page breaks

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    start_pos = len(text)
                    paragraph_text = paragraph.text + "\n\n"
                    text += paragraph_text
                    end_pos = len(text) - 1

                    estimated_page = max(1, (start_pos // chars_per_page) + 1)
                    page_mapping.append((start_pos, end_pos, estimated_page))

            # Extract from tables
            for table in doc.tables:
                start_pos = len(text)
                table_text = ""
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text += " | ".join(row_text) + "\n"
                table_text += "\n"
                text += table_text
                end_pos = len(text) - 1

                if table_text.strip():
                    estimated_page = max(1, (start_pos // chars_per_page) + 1)
                    page_mapping.append((start_pos, end_pos, estimated_page))

            return text, page_mapping

        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX file: {e}")

    def get_page_number_for_position(self, position: int, page_mapping: List[tuple]) -> int:
        """Get page number for a given text position."""
        for start_pos, end_pos, page_num in page_mapping:
            if start_pos <= position <= end_pos:
                return page_num
        # If not found, estimate based on closest page
        if page_mapping:
            for start_pos, end_pos, page_num in page_mapping:
                if position < start_pos:
                    return page_num
            return page_mapping[-1][2]
        return 1

    def chunk_text(self, text: str, page_mapping: List[tuple], chunk_size: int = 512,
                   similarity_threshold: float = 0.5) -> List:
        """
        Chunk text using Chonkie SemanticChunker with page tracking.
        Args:
            text: Input text to chunk
            page_mapping: List of (start_pos, end_pos, page_num) tuples
            chunk_size: Maximum tokens per chunk
            similarity_threshold: Similarity threshold for semantic chunking

        Returns:
            List of chunks with page number metadata
        """
        # Use model name string instead of SentenceTransformer object
        chunker = SemanticChunker(
            chunk_size=chunk_size,
            similarity_threshold=similarity_threshold,
            embedding_model=self.embedding_generator.model_name
        )
        chunks = chunker.chunk(text)

        # Add page number to each chunk
        for chunk in chunks:
            if hasattr(chunk, 'start_index') and page_mapping:
                chunk.page_number = self.get_page_number_for_position(chunk.start_index, page_mapping)
            else:
                chunk.page_number = 1

        return chunks

    async def process_document(self, file_path: str, chunk_size: int = 512,
                         similarity_threshold: float = 0.5,
                         document_id: str = None, metadata: Dict = None) -> str:
        """
        Process a document: extract text, chunk, embed, and store.

        Args:
            file_path: Path to the document file
            chunk_size: Maximum tokens per chunk
            similarity_threshold: Similarity threshold for chunking
            document_id: Optional document ID (if None, will generate one)
            metadata: Additional metadata for the document

        Returns:
            Document ID
        """
        file_path = Path(file_path)
        filename = file_path.name
        file_type = file_path.suffix.lower().replace('.', '')
        file_size = file_path.stat().st_size

        # Generate document ID if not provided
        if document_id is None:
            document_id = str(uuid.uuid4())

        print(f"Processing document: {filename} (ID: {document_id})")

        # Extract text based on file type
        page_mapping = None
        if file_type == 'pdf':
            text, page_mapping = self.extract_text_from_pdf(str(file_path))
        elif file_type == 'docx':
            text, page_mapping = self.extract_text_from_docx(str(file_path))
        elif file_type in ['txt']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            # For TXT files, create simple page mapping
            page_mapping = [(0, len(text) - 1, 1)] if text else []
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        print(f"Extracted {len(text)} characters from {filename}")

        # Chunk the text with page tracking
        chunks = self.chunk_text(text, page_mapping, chunk_size, similarity_threshold)
        print(f"Created {len(chunks)} chunks")

        # Prepare chunks for embedding
        chunk_texts = [chunk.text for chunk in chunks]

        # Generate embeddings in batch
        print("Generating embeddings...")
        embeddings = self.embedding_generator.embed_batch(chunk_texts)

        # Create Chunk objects using your interface
        chunk_objects = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_metadata = {
                'chunk_index': i,
                'token_count': chunk.token_count,
                'start_index': getattr(chunk, 'start_index', None),
                'end_index': getattr(chunk, 'end_index', None),
                'page_number': getattr(chunk, 'page_number', 1),
                'chunk_size': chunk_size,
                'similarity_threshold': similarity_threshold,
                'embedding_model': self.embedding_generator.model_name,
                'embedding_dimension': len(embedding),
                'filename': filename,
                'file_type': file_type,
                'file_size': file_size
            }

            # Add any additional metadata
            if metadata:
                chunk_metadata.update(metadata)

            chunk_obj = Chunk(
                id=str(uuid.uuid4()),
                document_id=document_id,
                text=chunk.text,
                embedding=embedding,
                metadata=chunk_metadata
            )
            chunk_objects.append(chunk_obj)

        # Use your add_chunks method with pgvector
        print("Inserting chunks into database using pgvector...")
        await self.vector_store.add_chunks(chunk_objects)

        print(
            f"Successfully processed {filename} -> Document ID: {document_id}")
        return document_id

    async def search_documents(self, query: str, limit: int = 5, threshold: float = 0.7,
                         document_ids: Optional[List[str]] = None) -> List[Dict]:
        """
        Search for relevant document chunks using pgvector.
        Args:
            query: Search query
            limit: Maximum number of results
            threshold: Similarity threshold
            document_ids: Optional list of document IDs to filter by

        Returns:
            List of relevant chunks
        """
        query_embedding = self.embedding_generator.embed_text(query)
        return await self.vector_store.search_similar_chunks(
            query_embedding, limit, threshold, document_ids
        )

    async def delete_document(self, document_id: str) -> int:
        """Delete all chunks for a document."""
        return await self.vector_store.delete_document_chunks(document_id)

    async def get_stats(self) -> Dict[str, Any]:
        """Get vector store statistics."""
        return await self.vector_store.get_collection_stats()
