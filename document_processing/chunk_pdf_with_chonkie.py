import PyPDF2
from docx import Document
from chonkie import SemanticChunker


def extract_text_from_pdf(pdf_path):
    """
    Extract text from PDF file using PyPDF2 with page tracking.
    Args:
        pdf_path (str): Path to the PDF file
    Returns:
        tuple: (full_text, page_mapping) where page_mapping is list of (start_pos, end_pos, page_num)
    """
    text = ""
    page_mapping = []

    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)

        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()

            start_pos = len(text)
            text += page_text + "\n"
            end_pos = len(text) - 1  # Exclude the newline

            if page_text.strip():  # Only add mapping if page has content
                page_mapping.append((start_pos, end_pos, page_num + 1))  # 1-indexed page numbers

    return text, page_mapping


def extract_text_from_docx(docx_path):
    """
    Extract text from DOCX file using python-docx with page estimation.
    Args:
        docx_path (str): Path to the DOCX file
    Returns:
        tuple: (full_text, page_mapping) where page_mapping estimates pages based on content
    """
    try:
        doc = Document(docx_path)
        text = ""
        page_mapping = []
        current_page = 1
        chars_per_page = 2500  # Rough estimate for page breaks

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            start_pos = len(text)
            paragraph_text = paragraph.text + "\n"
            text += paragraph_text
            end_pos = len(text) - 1

            if paragraph.text.strip():  # Only track non-empty paragraphs
                # Estimate page number based on character position
                estimated_page = max(1, (start_pos // chars_per_page) + 1)
                page_mapping.append((start_pos, end_pos, estimated_page))

        # Extract text from tables
        for table in doc.tables:
            start_pos = len(text)
            table_text = ""
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
                table_text += "\n"

            text += table_text
            end_pos = len(text) - 1

            if table_text.strip():
                estimated_page = max(1, (start_pos // chars_per_page) + 1)
                page_mapping.append((start_pos, end_pos, estimated_page))

        return text, page_mapping
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX file: {e}")


def chunk_with_semantic_chunker(text, chunk_size=512, similarity_threshold=0.5, embedding_model=None):
    """
    Chunk text using Chonkie's SemanticChunker with custom embedding model.
    Args:
        text (str): Text to chunk
        chunk_size (int): Maximum tokens per chunk
        similarity_threshold (float): Similarity threshold for semantic chunking
        embedding_model: Custom embedding model (SentenceTransformer or similar)
    Returns:
        list: List of chunks
    """
    if embedding_model:
        chunker = SemanticChunker(
            chunk_size=chunk_size,
            similarity_threshold=similarity_threshold,
            embedding_model=embedding_model
        )
    else:
        # Use default embedding model
        chunker = SemanticChunker(
            chunk_size=chunk_size,
            similarity_threshold=similarity_threshold
        )

    chunks = chunker.chunk(text)
    return chunks


def save_chunks_to_file(chunks, output_path, chunker_type):
    """
    Save chunks to a text file with page numbers.

    Args:
        chunks (list): List of chunks
        output_path (str): Output file path
        chunker_type (str): Type of chunker used
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(f"Chunks created using {chunker_type}\n")
        f.write("=" * 50 + "\n\n")

        for i, chunk in enumerate(chunks, 1):
            f.write(f"Chunk {i}:\n")
            f.write("-" * 20 + "\n")
            f.write(f"{chunk.text}\n\n")
            f.write(f"Tokens: {chunk.token_count}\n")
            if hasattr(chunk, 'page_number'):
                f.write(f"Page Number: {chunk.page_number}\n")
            if hasattr(chunk, 'start_index'):
                f.write(f"Start Index: {chunk.start_index}\n")
            if hasattr(chunk, 'end_index'):
                f.write(f"End Index: {chunk.end_index}\n")
            f.write("\n" + "="*50 + "\n\n")


def get_page_number_for_position(position, page_mapping):
    """
    Get page number for a given text position.
    Args:
        position (int): Character position in text
        page_mapping (list): List of (start_pos, end_pos, page_num) tuples
    Returns:
        int: Page number
    """
    for start_pos, end_pos, page_num in page_mapping:
        if start_pos <= position <= end_pos:
            return page_num
    # If not found, estimate based on closest page
    if page_mapping:
        for start_pos, end_pos, page_num in page_mapping:
            if position < start_pos:
                return page_num
        # If position is after all mapped content, return last page
        return page_mapping[-1][2]
    return 1  # Default to page 1

def process_document(file_path, chunk_size=512, similarity_threshold=0.5, embedding_model=None):
    """
    Process a document (PDF, DOCX, or TXT) and return chunks with page numbers.
    Args:
        file_path (str): Path to the document file
        chunk_size (int): Maximum tokens per chunk
        similarity_threshold (float): Similarity threshold for semantic chunking
        embedding_model: Custom embedding model
    Returns:
        tuple: (chunks, page_mapping) where chunks have page number metadata
    """
    from pathlib import Path

    file_path = Path(file_path)
    file_type = file_path.suffix.lower().replace('.', '')

    print(f"Processing {file_type.upper()} file: {file_path.name}")

    # Extract text based on file type
    page_mapping = None
    if file_type == 'pdf':
        text, page_mapping = extract_text_from_pdf(str(file_path))
    elif file_type == 'docx':
        text, page_mapping = extract_text_from_docx(str(file_path))
    elif file_type == 'txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        # For TXT files, create a simple page mapping (assume one page)
        page_mapping = [(0, len(text) - 1, 1)] if text else []
    else:
        raise ValueError(
            f"Unsupported file type: {file_type}. Supported types: PDF, DOCX, TXT")

    print(f"Extracted {len(text)} characters from {file_path.name}")

    # Chunk the text
    chunks = chunk_with_semantic_chunker(
        text, chunk_size, similarity_threshold, embedding_model)
    print(f"Created {len(chunks)} chunks")

    # Add page number metadata to chunks
    for chunk in chunks:
        if hasattr(chunk, 'start_index') and page_mapping:
            chunk.page_number = get_page_number_for_position(chunk.start_index, page_mapping)
        else:
            chunk.page_number = 1  # Default to page 1 if no position info

    return chunks, page_mapping
